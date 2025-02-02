from base64 import b64decode
from bs4 import BeautifulSoup
from collections import OrderedDict
from io import BytesIO
from logging import getLogger

from docx import Document
from docxcompose.composer import Composer
from docxtpl import DocxTemplate
from requests import codes as codes_request, post as post_request
from requests.exceptions import RequestException

from odoo import _, api, fields, models
from odoo.exceptions import AccessError, UserError
from odoo.http import request
from odoo.tools.safe_eval import safe_eval, time

try:
    from odoo.addons.gotenberg.service.utils import (
        get_auth,  # noqa
        convert_pdf_from_office_url,  # noqa
        check_gotenberg_installed,  # noqa
    )

    gotenberg_imported = True
except ImportError:
    gotenberg_imported = False

_logger = getLogger(__name__)


class IrActionsReport(models.Model):
    _inherit = "ir.actions.report"

    report_name = fields.Char(
        compute="_compute_report_name",
        inverse="_inverse_report_name",
        store=True,
        required=False,
    )
    report_type = fields.Selection(
        selection_add=[("docx-docx", "DOCX"), ("docx-pdf", "DOCX(PDF)")],
        ondelete={"docx-docx": "cascade", "docx-pdf": "cascade"},
    )
    report_docx_template = fields.Binary(
        string="Report docx template",
    )

    @api.depends("report_type", "model")
    def _compute_report_name(self):
        for record in self:
            if (
                record.report_type in ["docx-docx", "docx-pdf"]
                and record.model
                and record.id
            ):
                record.report_name = "%s-docx_report+%s" % (record.model, record.id)
            else:
                record.report_name = False

    def _inverse_report_name(self):
        """TODO: write this method"""
        pass

    def retrieve_attachment(self, record):
        """
        Searc for existing report file in record's attachments by fields:
        1. name
        2. res_model
        3. res_id
        """
        result = super().retrieve_attachment(record)
        if result:
            if self.report_type == "docx-docx":
                result = (
                    result.filtered(
                        lambda r: r.mimetype
                        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    or None
                )
            elif self.report_type == "docx-pdf":
                result = (
                    result.filtered(lambda r: r.mimetype == "application/pdf") or None
                )
        return result

    @api.model
    def _render_docx_pdf(self, res_ids=None, data=None):
        """
        Prepares the data for report file rendering, calls for the render method
        and handle rendering result.
        """
        if not data:
            data = {}
        data.setdefault("report_type", "pdf")

        # access the report details with sudo() but evaluation context as current user
        self_sudo = self.sudo()

        save_in_attachment = OrderedDict()
        # Maps the streams in `save_in_attachment` back to the records they came from
        # stream_record = dict()
        if res_ids:
            Model = self.env[self_sudo.model]
            record_ids = Model.browse(res_ids)
            docx_record_ids = Model
            if self_sudo.attachment:
                for record_id in record_ids:
                    attachment = self_sudo.retrieve_attachment(record_id)
                    if attachment and self_sudo.attachment_use:
                        # stream = self_sudo._retrieve_stream_from_attachment(attachment)
                        stream = BytesIO(attachment.raw)
                        save_in_attachment[record_id.id] = stream
                        # stream_record[stream] = record_id
                    if not self_sudo.attachment_use or not attachment:
                        docx_record_ids += record_id
            else:
                docx_record_ids = record_ids
            res_ids = docx_record_ids.ids

        if save_in_attachment:  # and not res_ids:
            _logger.info("The PDF report has been generated from attachment.")
            # self._raise_on_unreadable_pdfs(save_in_attachment.values(), stream_record)
            return self_sudo._post_pdf(save_in_attachment), "pdf"

        docx_content = self._render_docx(res_ids, data=data)

        pdf_content = (
            self._get_pdf_from_office(docx_content)
            if gotenberg_imported and check_gotenberg_installed()
            else None
        )

        if not pdf_content:
            raise UserError(
                _(
                    "Gotenberg converting service not available. The PDF can not be created."
                )
            )

        if res_ids:
            # self._raise_on_unreadable_pdfs(save_in_attachment.values(), stream_record)
            # saving pdf in attachment.
            return (
                self_sudo._post_pdf(
                    save_in_attachment, pdf_content=pdf_content, res_ids=res_ids
                ),
                "pdf",
            )

        return pdf_content, "pdf"

    @api.model
    def _render_docx_docx(self, res_ids=None, data=None):
        """
        Prepares the data for report file rendering, calls for the render method
        and handle rendering result.
        """
        if not data:
            data = {}
        data.setdefault("report_type", "docx")

        # access the report details with sudo() but evaluation context as current user
        self_sudo = self.sudo()

        save_in_attachment = OrderedDict()
        # Maps the streams in `save_in_attachment` back to the records they came from
        stream_record = dict()
        if res_ids:
            Model = self.env[self_sudo.model]
            record_ids = Model.browse(res_ids)
            docx_record_ids = Model
            if self_sudo.attachment:
                for record_id in record_ids:
                    attachment = self_sudo.retrieve_attachment(record_id)
                    if attachment:
                        # stream = self_sudo._retrieve_stream_from_attachment(attachment)
                        stream = BytesIO(attachment.raw)
                        save_in_attachment[record_id.id] = stream
                        stream_record[stream] = record_id
                    if not self_sudo.attachment_use or not attachment:
                        docx_record_ids += record_id
            else:
                docx_record_ids = record_ids
            res_ids = docx_record_ids.ids

        if save_in_attachment and not res_ids:
            _logger.info("The DOCS report has been generated from attachment.")
            return self_sudo._post_docx(save_in_attachment), "docx"

        docx_contents = []
        for record_id in res_ids:
            docx_content = self._render_docx([record_id], data=data)
            docx_contents.append(docx_content)

        if res_ids:
            _logger.info(
                "The DOCS report has been generated for model: %s, records %s."
                % (self_sudo.model, str(res_ids))
            )
            return (
                self_sudo._post_docx(
                    save_in_attachment, docx_contents=docx_contents, res_ids=res_ids
                ),
                "docx",
            )
        return docx_contents, "docx"

    def _post_pdf(self, save_in_attachment, pdf_content=None, res_ids=None):
        """
        Adds pdf file in record's attachments.
        TODO: For now bunch generation is not supported.
        2 execution ways:
           - save_in_attachment and not res_ids - when get reports from attachments
           - res_ids and not save_in_attachment - when generate report.
        """
        self_sudo = self.sudo()
        attachment_vals_list = []
        if save_in_attachment:
            # here get streams from save_in_attachment, make pdf file and return it
            # bunch generation here is already realized.
            reports_data = list(save_in_attachment.values())
            if len(reports_data) == 1:
                # If only one report, no need to merge files. Returns as is.
                return reports_data[0].getvalue()
            else:
                return self._merge_pdfs(reports_data)
        for res_id in res_ids:
            record = self.env[self_sudo.model].browse(res_id)
            if not self_sudo.attachment:
                attachment_name = False
            else:
                attachment_name = safe_eval(
                    self_sudo.attachment, {"object": record, "time": time}
                )
            # Unable to compute a name for the attachment.
            if not attachment_name:
                continue
            attachment_vals_list.append(
                {
                    "name": attachment_name,
                    "raw": pdf_content,  # stream_data['stream'].getvalue(),
                    "res_model": self_sudo.model,
                    "res_id": record.id,
                    "type": "binary",
                }
            )
        if attachment_vals_list:
            attachment_names = ", ".join(x["name"] for x in attachment_vals_list)
            try:
                self.env["ir.attachment"].create(attachment_vals_list)
            except AccessError:
                _logger.info(
                    "Cannot save PDF report %r attachments for user %r",
                    attachment_names,
                    self.env.user.display_name,
                )
            else:
                _logger.info(
                    "The PDF documents %r are now saved in the database",
                    attachment_names,
                )
        return pdf_content

    def _post_docx(self, save_in_attachment, docx_contents=None, res_ids=None):
        """
        Adds generated file in attachments.
        """

        def close_streams(streams):
            for stream in streams:
                try:
                    stream.close()
                except Exception:
                    pass

        if len(save_in_attachment) == 1 and not docx_contents:
            return list(save_in_attachment.values())[0].getvalue()

        streams = []
        if docx_contents:
            # Build a record_map mapping id -> record
            record_map = {
                r.id: r
                for r in self.env[self.model].browse(
                    [res_id for res_id in res_ids if res_id]
                )
            }
            # If no value in attachment or no record specified, only append the whole docx.
            if not record_map or not self.attachment:
                streams.extend(docx_contents)
            else:
                for res_id, docx_content in zip(res_ids, docx_contents):
                    if res_id in record_map and not res_id in save_in_attachment:
                        new_stream = self._postprocess_docx_report(
                            record_map[res_id], docx_content
                        )
                        # If the buffer has been modified, mark the old buffer to be closed as well.
                        if new_stream and new_stream != docx_content:
                            close_streams([docx_content])
                            docx_content = new_stream
                    streams.append(docx_content)
        if self.attachment_use:
            for stream in save_in_attachment.values():
                streams.append(stream)
        if len(streams) == 1:
            result = streams[0].getvalue()
        else:
            try:
                merged_stream = self._merge_docx(streams)
                result = merged_stream.getvalue()
            except Exception as e:
                _logger.exception(e)
                raise UserError(
                    _("One of the documents you try to merge caused failure.")
                )

        close_streams(streams)
        return result

    def _postprocess_docx_report(self, record, buffer):
        """
        Creates the record in the "ir.attachment" model.
        """
        attachment_name = safe_eval(self.attachment, {"object": record, "time": time})
        if not attachment_name:
            return None
        attachment_vals = {
            "name": attachment_name,
            "raw": buffer.getvalue(),
            "res_model": self.model,
            "res_id": record.id,
            "type": "binary",
        }
        try:
            self.env["ir.attachment"].create(attachment_vals)
        except AccessError:
            _logger.info(
                "Cannot save DOCX report %r as attachment", attachment_vals["name"]
            )
        else:
            _logger.info(
                "The DOCX document %s is now saved in the database",
                attachment_vals["name"],
            )
        return buffer

    @staticmethod
    def _merge_docx(streams):
        """
        Joins several docx files into one with page breaks between them.
        """
        if not streams:
            return None

        merged_document = Document()
        composer = Composer(merged_document)

        for stream in streams:
            document = Document(stream)

            if composer.doc.paragraphs:
                composer.doc.add_page_break()

            composer.append(document)

        merged_stream = BytesIO()
        merged_document.save(merged_stream)
        merged_stream.seek(0)

        return merged_stream

    def _render_docx(self, docids: list, data: dict = None):
        """
        Receive the data for rendering and calls for it.

        docids: list of record's ids for which report is generated.
        data: dict, conains "context", "report_type".
        """
        if not data:
            data = {}
        data.setdefault("report_type", "docx")
        data = self._get_rendering_context(
            self, docids, data
        )  # self contains current record of ir.actions.report model.
        return self._render_docx_template(self.report_docx_template, values=data)

    def _render_docx_template(self, template: bytes, values: dict = None):
        """
        docx file rendering itself.
        """
        if values is None:
            values = {}
        context = dict(self.env.context, inherit_branding=False)
        # Browse the user instead of using the sudo self.env.user
        user = self.env["res.users"].browse(self.env.uid)
        website = None
        if request and hasattr(request, "website"):
            if request.website is not None:
                website = request.website
                context = dict(
                    context,
                    translatable=context.get("lang")
                    != request.env["ir.http"]._get_default_lang().code,
                )
        values.update(
            record=values["docs"],
            time=time,
            context_timestamp=lambda t: fields.Datetime.context_timestamp(
                self.with_context(tz=user.tz), t
            ),
            user=user,
            res_company=user.company_id,
            website=website,
            web_base_url=self.env["ir.config_parameter"]
            .sudo()
            .get_param("web.base.url", default=""),
        )

        record_to_render = values["docs"]
        docs = {
            key: record_to_render[key]
            for key in record_to_render._fields.keys()
            if not isinstance(record_to_render[key], fields.Markup)
        }
        docs.update(
            {
                key: self._parse_markup(record_to_render[key])
                for key in record_to_render._fields.keys()
                if isinstance(record_to_render[key], fields.Markup)
            }
        )
        values["docs"] = docs

        docx_content = BytesIO()
        with BytesIO(b64decode(template)) as template_file:
            doc = DocxTemplate(template_file)
            doc.render(values)
            doc.save(docx_content)
        docx_content.seek(0)
        return docx_content

    @staticmethod
    def _parse_markup(markup_data: fields.Markup):
        """
        Extracts data from field of Html type and returns them in text format,
        without html tags.
        """
        soup = BeautifulSoup(markup_data.__str__())
        data_arr = list(soup.strings)
        return "\n".join(data_arr)

    @staticmethod
    def _get_pdf_from_office(content_stream):
        """
        Converting docx into pdf with Gotenberg service.
        """
        result = None
        url = convert_pdf_from_office_url()
        auth = get_auth()
        try:
            response = post_request(
                url,
                files={"file": ("converted_file.docx", content_stream.read())},
                auth=auth,
            )
            if response.status_code == codes_request.ok:
                result = response.content
            else:
                _logger.warning(
                    "Gotenberg response: %s - %s"
                    % (response.status_code, response.content)
                )
        except RequestException as e:
            _logger.exception(e)
        finally:
            return result
